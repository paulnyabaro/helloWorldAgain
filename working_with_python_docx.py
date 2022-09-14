import docx
doc = docx.Document('example_doc.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

# Indentation in word documents
# fullText.append('    ' + para.text)

# Double space
# return '\n\n'.join(fullText)

# Default styles
# 'Normal'
# 'BodyText'
# 'BodyText2'
# 'BodyText3'
# 'Caption'
# 'Heading1'
# 'Heading2'
# 'Heading3'
# 'Heading4'
# 'Heading5'
# 'Heading6'
# 'Heading7'
# 'Heading8'
# 'Heading9'
# 'IntenseQuote'
# 'List'
# 'List2'
# 'List3'
# 'ListBullet'
# 'ListBullet2'
# 'ListBullet3'
# 'ListContinue'
# 'ListContinue2'
# 'ListContinue3'
# 'ListNumber'
# 'ListNumber2'
# 'ListNumber3'
# 'ListParagraph'
# 'MacroText'
# 'NoSpacing'
# 'Quote'
# 'Subtitle'
# 'TOCHeading'
# 'Title'

# Applying styling
# doc.paragraphs[1].runs[1].underline = True