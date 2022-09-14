import docx
doc = docx.Document()
doc.add_paragraph('Hello world again!')

# Adding other paragraphs
paraObj1 = doc.add_paragraph('This is a second paragraph.', 'Title') # Adding title
# paraObj1 = doc.add_paragraph('This is a second paragraph.', '0') # 0 is title 1 - 4 heading levels
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('helloworld.docx')

# Adding page breaks
# doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)

# Adding pictures
# doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))